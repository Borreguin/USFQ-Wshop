from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/mnt/data/tsp_ga_project_v4')
REPORT = ROOT / 'report'
FIG = ROOT / 'figures'
RES = ROOT / 'results'
REPORT.mkdir(exist_ok=True)

summary = pd.read_csv(RES / 'summary_all_datasets.csv')
validation = pd.read_csv(RES / 'validacion_permutaciones.csv')
params = pd.read_csv(RES / 'parameter_experiment_largest_dataset.csv')

avg_gap_nn = summary['gap_vs_nn_pct'].mean()
avg_gap_nn2 = summary['gap_vs_nn2_pct'].mean()
avg_nn = summary['nn_distance'].mean()
avg_ga = summary['ga_2opt_distance'].mean()
valid_all = summary['valid_solution'].all()
premature_count = summary['stagnation'].str.contains('prematura', case=False, na=False).sum()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=8.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)


def add_table_from_df(doc, df, columns, headers=None, font_size=8.0):
    headers = headers or columns
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, font_size=font_size)
        set_cell_shading(hdr[i], 'D9EAF7')
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                val = f'{val:.2f}'
            set_cell_text(cells[i], val, font_size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_note(doc, title, text):
    p = doc.add_paragraph()
    p.style = 'Intense Quote'
    r = p.add_run(title + ' ')
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(text)


def add_image(doc, path, caption, width=6.4):
    if Path(path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)


def round_df(df, cols):
    out = df.copy()
    for c in cols:
        out[c] = out[c].astype(float).round(2)
    return out


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal'].font.size = Pt(11)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.name = 'Times New Roman'
    styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
styles['Heading 1'].font.size = Pt(16)
styles['Heading 2'].font.size = Pt(13)
styles['Heading 3'].font.size = Pt(12)

# Header/footer
for sec in doc.sections:
    sec.header.paragraphs[0].text = 'Proyecto final de IA - TSP mediante Algoritmos Geneticos'
    sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sec.footer.paragraphs[0].text = 'Grupo 5'
    sec.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nPROYECTO FINAL DE INTELIGENCIA ARTIFICIAL\n')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Resolucion del TSP mediante Algoritmos Geneticos\n')
run.bold = True
run.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Informe profesional del proyecto\n')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\nGrupo 5\n').bold = True
for name in ['Nancy Altamirano', 'Gustavo Berru', 'Raquel Pacheco', 'Kevin Viteri']:
    p = doc.add_paragraph(name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\nCiencia de Datos\n').bold = True
p.add_run('Mayo de 2026')

doc.add_page_break()

# Index
h = doc.add_heading('Indice del informe', level=1)
for item in [
    'Resumen ejecutivo',
    'Objetivo y enfoque metodologico',
    'a) Representacion de individuos como permutaciones validas',
    'b) Operadores geneticos seleccionados y justificacion',
    'c) Analisis de la evolucion del fitness',
    'd) Medicion y analisis de diversidad poblacional',
    'e) Evaluacion de la calidad de las soluciones',
    'f) Analisis experimental de parametros',
    'Conclusiones tecnicas',
    'Resumen de participacion del Grupo 5',
    'Referencias de apoyo'
]:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Executive summary
h = doc.add_heading('Resumen ejecutivo', level=1)
doc.add_paragraph(
    'Este proyecto resuelve el Problema del Viajante de Comercio (TSP) mediante un Algoritmo Genetico Generacional Elitista. '
    'El TSP exige encontrar un recorrido cerrado que visite cada ciudad exactamente una vez y retorne al origen minimizando la distancia total. '
    'Por esa razon, la solucion no se represento como cadena binaria, sino como una permutacion de indices de ciudades. Esta decision evita, desde el diseno del cromosoma, rutas con ciudades repetidas o faltantes.'
)
doc.add_paragraph(
    f'Se procesaron 10 datasets con entre {summary["n_ciudades"].min()} y {summary["n_ciudades"].max()} ciudades. '
    'El algoritmo uso seleccion por torneo, cruce OX, mutacion por inversion, elitismo de dos individuos y mejora final 2-opt. '
    f'Todas las soluciones finales fueron validas ({valid_all}) y no se generaron hijos invalidos durante la ejecucion. '
    f'En promedio, el AG elitista + 2-opt redujo la distancia frente a Nearest Neighbor en {abs(avg_gap_nn):.2f}%. '
    f'Frente a Nearest Neighbor con 2-opt, el promedio fue {avg_gap_nn2:.2f}%, lo que muestra que el AG fue competitivo incluso contra una heuristica ya refinada.'
)
add_note(doc, 'Idea central para la sustentación:',
         'El algoritmo elegido es un Algoritmo Genetico Generacional Elitista porque conserva las mejores rutas entre generaciones y reduce el riesgo de perder buenas soluciones por azar. OX no es otro algoritmo ni reemplaza al elitismo: OX es el operador de crossover incorporado dentro del AG para crear descendientes validos cuando las soluciones son permutaciones de ciudades.')

# Objective
h = doc.add_heading('Objetivo y enfoque metodologico', level=1)
doc.add_paragraph(
    'El objetivo del ejercicio es implementar y analizar una solucion evolutiva para el TSP cumpliendo seis literales: representacion valida, operadores geneticos adecuados, evolucion del fitness, diversidad poblacional, comparacion con una heuristica y analisis experimental de parametros. Para sustentar visualmente los resultados, el informe incluye graficas de fitness, convergencia, diversidad, comparacion con heuristicas y analisis de parametros. '
    'El enfoque utilizado fue experimental y reproducible: los datasets se leyeron desde archivos CSV con columnas city, x e y; luego se calculo una matriz de distancias euclidianas y se ejecuto el algoritmo con semilla fija para permitir replicacion.'
)

doc.add_paragraph('Configuracion base usada en la ejecucion principal:')
config_df = pd.DataFrame([
    ['Representacion', 'Permutacion de ciudades'],
    ['Fitness', '1 / distancia_total'],
    ['Tamano poblacional', 100],
    ['Generaciones', 250],
    ['Seleccion', 'Torneo k = 3'],
    ['Crossover', 'OX, probabilidad 0.90'],
    ['Mutacion', 'Inversion, probabilidad 0.15'],
    ['Elitismo', '2 mejores individuos preservados'],
    ['Comparacion', 'Nearest Neighbor y Nearest Neighbor + 2-opt'],
    ['Mejora local final', '2-opt aplicado a la mejor ruta del AG']
], columns=['Elemento', 'Valor'])
add_table_from_df(doc, config_df, ['Elemento', 'Valor'], font_size=9)

# Dataset table
h = doc.add_heading('Datasets trabajados', level=2)
datasets_df = summary[['dataset', 'n_ciudades']].copy()
add_table_from_df(doc, datasets_df, ['dataset', 'n_ciudades'], ['Dataset', 'Numero de ciudades'], font_size=8.5)

# a
h = doc.add_heading('a) Representacion de cada individuo como una permutacion valida de ciudades', level=1)
doc.add_paragraph(
    'Cada individuo se represento como un vector ordenado de indices de ciudades. Si existen n ciudades, un cromosoma valido contiene exactamente los valores de 0 a n-1, sin repetir ninguno. '
    'El orden de los genes define el orden de visita. La ruta se considera cerrada porque, despues de visitar la ultima ciudad, el algoritmo suma la distancia de retorno hacia la primera ciudad.'
)
doc.add_paragraph(
    'Ejemplo conceptual: [4, 1, 7, 0, 3, 2, 6, 5] significa que el viajante inicia en la ciudad 4, continua por la 1, luego por la 7, y asi sucesivamente, hasta volver a la ciudad 4. '
    'Esta representacion es la adecuada para el TSP porque el problema no busca activar o desactivar ciudades, sino ordenar todas las ciudades disponibles.'
)

doc.add_paragraph('Garantia de validez solicitada en el literal:')
validity_df = pd.DataFrame([
    ['Inicializacion', 'Se genera cada individuo con una permutacion aleatoria de 0..n-1. Tambien se incluye una ruta semilla Nearest Neighbor, que igualmente contiene todas las ciudades una sola vez.', 'Sin repetidas y sin faltantes desde la primera poblacion.'],
    ['Crossover OX', 'Copia un segmento del padre 1 y completa con las ciudades del padre 2 que aun no estan en el hijo.', 'No duplica ciudades porque solo inserta genes no usados; no omite ciudades porque completa todas las posiciones.'],
    ['Mutacion por inversion', 'Invierte el orden de un segmento interno de la misma ruta.', 'No cambia el conjunto de ciudades; solo cambia su posicion.'],
    ['Validacion en codigo', 'Se usa is_valid_permutation(route, n) y validate_population(population, n).', 'Se confirma longitud n, conjunto completo y ausencia de repetidos.']
], columns=['Fase', 'Mecanismo aplicado', 'Justificacion de validez'])
add_table_from_df(doc, validity_df, ['Fase', 'Mecanismo aplicado', 'Justificacion de validez'], font_size=7.6)

doc.add_paragraph(
    'La tabla de validacion confirma que todos los datasets terminaron con rutas validas. Ademas, la ejecucion registro cero hijos invalidos, cero hijos reparados y cero generaciones con poblaciones invalidas. '
    'Esto es importante porque no basta con encontrar una distancia baja: la solucion debe ser factible para el TSP.'
)
val_table = validation[['dataset', 'n_ciudades', 'ruta_ag_2opt_valida', 'hijos_invalidos_generados', 'generaciones_con_poblacion_invalida', 'sin_ciudades_repetidas', 'sin_ciudades_faltantes']].copy()
add_table_from_df(doc, val_table, ['dataset', 'n_ciudades', 'ruta_ag_2opt_valida', 'hijos_invalidos_generados', 'generaciones_con_poblacion_invalida', 'sin_ciudades_repetidas', 'sin_ciudades_faltantes'],
                  ['Dataset', 'n', 'Ruta final valida', 'Hijos invalidos', 'Poblaciones invalidas', 'Sin repetidas', 'Sin faltantes'], font_size=7.2)
add_note(doc, 'Sustentación del literal a):',
         'La validez no se deja al azar. La representacion, el crossover y la mutacion fueron elegidos para trabajar naturalmente con permutaciones; ademas, el codigo valida cada ruta final y registra si aparecieron soluciones invalidas.')

# b
h = doc.add_heading('b) Definicion de operadores geneticos y justificacion del uso', level=1)
doc.add_paragraph(
    'El algoritmo principal es un AG generacional elitista. Esto significa que en cada generacion se crea una nueva poblacion de descendientes, pero los mejores individuos de la generacion actual pasan directamente a la siguiente. '
    'El elitismo se eligio porque en el TSP una buena ruta puede perderse por azar durante el cruce o la mutacion; conservarla evita retrocesos en la mejor solucion encontrada.'
)
ops_df = pd.DataFrame([
    ['Seleccion por torneo', 'Se eligen k=3 candidatos al azar y gana el de menor distancia.', 'Controla la presion selectiva: favorece rutas buenas sin eliminar totalmente la diversidad.'],
    ['Crossover OX', 'Se preserva un segmento de un padre y se completa el resto siguiendo el orden del otro padre.', 'Es valido para permutaciones y conserva informacion de orden, que es clave en el TSP.'],
    ['Mutacion por inversion', 'Se seleccionan dos posiciones y se invierte el segmento entre ellas.', 'Introduce variacion sin crear ciudades repetidas ni faltantes. Ademas, es coherente con la idea de mejorar tramos de una ruta.'],
    ['Elitismo', 'Los 2 mejores individuos se copian directamente a la siguiente generacion.', 'Evita perder la mejor ruta encontrada y mejora la estabilidad de la convergencia.'],
    ['2-opt final', 'Se invierten aristas si la inversion reduce la distancia.', 'Refina localmente la mejor ruta del AG; por eso el enfoque puede sustentarse como AG elitista hibrido.']
], columns=['Componente', 'Como funciona', 'Justificacion'])
add_table_from_df(doc, ops_df, ['Componente', 'Como funciona', 'Justificacion'], font_size=7.6)

doc.add_paragraph(
    'Se eligio OX como operador de crossover porque el TSP depende del orden de visita. Un cruce tradicional de cadenas binarias podria generar rutas con ciudades repetidas o faltantes. '
    'OX evita ese problema: primero fija un bloque del padre 1 y luego recorre el padre 2 para insertar solamente las ciudades que todavia no aparecen en el hijo. '
    'La mutacion por inversion se eligio porque cambia la geometria de la ruta sin alterar el conjunto de ciudades; por tanto, es mas apropiada para rutas que una mutacion numerica comun.'
)
add_note(doc, 'Sustentación del literal b):',
         'No se debe interpretar OX como una alternativa al AG elitista. El AG elitista es la variante principal del proceso evolutivo porque define como se conserva la mejor solucion entre generaciones; OX, en cambio, es el operador de crossover usado dentro de ese algoritmo para recombinar dos rutas padres sin romper la representacion por permutaciones. Por eso, la decision metodologica completa es: usar un AG generacional elitista y, dentro de el, aplicar OX como operador de cruce valido para el TSP.')

# c
h = doc.add_heading('c) Analisis de la evolucion del fitness', level=1)
doc.add_paragraph(
    'El fitness se definio como el inverso de la distancia total: fitness = 1 / distancia_total. Con esta definicion, una ruta mas corta tiene mayor fitness. '
    'En cada generacion se registraron tres valores: mejor fitness, fitness promedio y peor fitness. Tambien se guardaron las distancias equivalentes porque son mas faciles de interpretar en terminos del objetivo real del TSP.'
)
add_image(doc, FIG / 'panel_fitness_all_datasets.png', 'Figura 1. Evolucion del mejor fitness, fitness promedio y peor fitness por dataset.', width=6.7)

doc.add_paragraph(
    'La figura muestra el comportamiento esperado de un AG: el mejor fitness tiende a aumentar y luego estabilizarse; el fitness promedio se acerca gradualmente al mejor fitness conforme la poblacion converge; y el peor fitness puede fluctuar por efecto de la mutacion, pero suele mejorar cuando la poblacion ya ha heredado buenos patrones de ruta. '
    'Cuando la diferencia entre el mejor y el promedio se reduce demasiado rapido, puede existir riesgo de convergencia prematura, porque muchos individuos se parecen entre si y la exploracion disminuye.'
)
conv_df = summary[['dataset', 'best_generation', 'final_no_improve', 'unique_pct_final', 'hamming_pct_final', 'stagnation']].copy()
conv_df['stagnation'] = conv_df['stagnation'].str.replace('Convergencia activa: ', '', regex=False).str.replace('Posible convergencia prematura: ', '', regex=False)
add_table_from_df(doc, conv_df, ['dataset', 'best_generation', 'final_no_improve', 'unique_pct_final', 'hamming_pct_final', 'stagnation'],
                  ['Dataset', 'Mejor gen.', 'Gen. sin mejora al final', 'Unicos final (%)', 'Hamming final (%)', 'Diagnostico'], font_size=6.9)

doc.add_paragraph(
    f'El analisis de estancamiento se realizo observando cuantas generaciones finales pasaron sin mejora y si la diversidad final era baja. '
    f'En {premature_count} de los 10 datasets se identifico posible convergencia prematura, porque hubo un periodo final largo sin mejora junto con baja diversidad. '
    'En los demas casos se observo convergencia activa, debido a que todavia existieron mejoras dentro del tramo final.'
)
add_image(doc, FIG / 'comparacion_nn_vs_ag2opt.png', 'Figura 2. Comparacion de distancia final entre Nearest Neighbor y AG elitista + 2-opt.', width=6.6)
add_note(doc, 'Sustentación del literal c):',
         'La evolucion del fitness se grafico por generacion para todos los datasets. El criterio de convergencia no se baso solo en mirar la grafica, sino tambien en medir generaciones sin mejora y diversidad final.')

# d
h = doc.add_heading('d) Medicion y analisis de diversidad poblacional', level=1)
doc.add_paragraph(
    'Se usaron dos metricas de diversidad. La primera fue el porcentaje de individuos unicos en la poblacion. La segunda fue la distancia Hamming promedio normalizada entre pares de individuos, que mide el porcentaje de posiciones en las que dos rutas difieren. '
    'Ambas metricas son apropiadas para permutaciones: el porcentaje de unicos detecta duplicacion poblacional y la distancia Hamming permite observar si las rutas tienen ordenamientos parecidos.'
)
metric_df = pd.DataFrame([
    ['Porcentaje de individuos unicos', 'Numero de rutas distintas / tamano de poblacion x 100', 'Bajo valor indica que muchos cromosomas se repiten.'],
    ['Distancia Hamming promedio', 'Promedio del porcentaje de posiciones diferentes entre pares de rutas', 'Bajo valor indica rutas muy parecidas y posible perdida de exploracion.']
], columns=['Metrica', 'Calculo', 'Interpretacion'])
add_table_from_df(doc, metric_df, ['Metrica', 'Calculo', 'Interpretacion'], font_size=8.3)
add_image(doc, FIG / 'panel_diversidad_all_datasets.png', 'Figura 3. Evolucion de diversidad poblacional por dataset.', width=6.7)

doc.add_paragraph(
    'La diversidad disminuye conforme el algoritmo avanza porque la seleccion y el elitismo favorecen la reproduccion de rutas de mejor calidad. Esto es normal y deseable hasta cierto punto: una diversidad muy alta al final puede indicar que la poblacion aun no explota buenas soluciones; pero una diversidad demasiado baja puede indicar convergencia prematura. '
    'Por eso se analizo la relacion diversidad-calidad y diversidad-convergencia en conjunto, no de manera aislada.'
)
div_df = summary[['dataset', 'unique_pct_final', 'hamming_pct_final', 'corr_unique_vs_best_distance', 'corr_unique_vs_no_improve']].copy()
add_table_from_df(doc, div_df, ['dataset', 'unique_pct_final', 'hamming_pct_final', 'corr_unique_vs_best_distance', 'corr_unique_vs_no_improve'],
                  ['Dataset', 'Unicos final (%)', 'Hamming final (%)', 'Corr. diversidad-distancia', 'Corr. diversidad-sin mejora'], font_size=7.4)

doc.add_paragraph(
    'La correlacion diversidad-distancia se interpreta con cuidado: a medida que baja la diversidad, la distancia suele bajar porque la poblacion converge hacia mejores rutas. Sin embargo, si la diversidad cae demasiado y el numero de generaciones sin mejora aumenta, la convergencia puede volverse prematura. '
    'En este proyecto, los datasets cities_100_222324 y cities_100_252627 mostraron el patron mas claro de riesgo: baja diversidad final y muchas generaciones sin mejora.'
)
add_image(doc, FIG / 'diversidad_final_por_dataset.png', 'Figura 4. Diversidad final medida por porcentaje de individuos unicos.', width=6.4)
add_note(doc, 'Sustentación del literal d):',
         'Se midio la diversidad con dos indicadores complementarios. Esto permite explicar no solo si el algoritmo encontro buenas rutas, sino tambien si las encontro manteniendo suficiente exploracion poblacional.')

# e
h = doc.add_heading('e) Evaluacion de la calidad de las soluciones', level=1)
doc.add_paragraph(
    'La calidad de las soluciones se evaluo comparando el AG con una heuristica simple: Nearest Neighbor. Esta heuristica construye una ruta desde cada ciudad inicial y siempre selecciona como siguiente destino la ciudad no visitada mas cercana. '
    'Luego se conserva la mejor ruta obtenida entre todos los puntos de partida. Tambien se calculo Nearest Neighbor + 2-opt para tener una referencia mas fuerte.'
)
doc.add_paragraph(
    'Como los datasets no incluyen el optimo teorico certificado, no se puede afirmar error relativo respecto al optimo global. Por transparencia, el error relativo se reporto respecto a las heuristicas de referencia. '
    'La formula usada fue: gap relativo (%) = ((distancia_AG - distancia_referencia) / distancia_referencia) x 100. Un valor negativo significa que el AG encontro una ruta mas corta que la referencia.'
)
quality_df = summary[['dataset', 'n_ciudades', 'nn_distance', 'nn_2opt_distance', 'ga_2opt_distance', 'gap_vs_nn_pct', 'gap_vs_nn2_pct', 'valid_solution']].copy()
add_table_from_df(doc, quality_df, ['dataset', 'n_ciudades', 'nn_distance', 'nn_2opt_distance', 'ga_2opt_distance', 'gap_vs_nn_pct', 'gap_vs_nn2_pct', 'valid_solution'],
                  ['Dataset', 'n', 'NN', 'NN+2opt', 'AG+2opt', 'Gap vs NN (%)', 'Gap vs NN+2opt (%)', 'Valida'], font_size=6.9)

doc.add_paragraph(
    f'En promedio, Nearest Neighbor obtuvo una distancia de {avg_nn:.2f}, mientras que AG elitista + 2-opt obtuvo {avg_ga:.2f}. '
    f'Esto equivale a una mejora promedio de {abs(avg_gap_nn):.2f}% frente a la heuristica simple. '
    f'Frente a NN+2-opt, el promedio fue {avg_gap_nn2:.2f}%; es decir, el AG fue competitivo incluso contra una solucion heuristica refinada. '
    'El caso cities_100_192021 tuvo un gap positivo frente a NN+2-opt, lo que significa que la heuristica refinada fue mejor en ese dataset especifico; esto es normal en algoritmos aproximados y demuestra por que es importante comparar y no asumir optimalidad.'
)
add_image(doc, FIG / 'gap_relativo_vs_nn.png', 'Figura 5. Gap relativo frente a Nearest Neighbor. Valores negativos representan mejora.', width=6.4)
add_note(doc, 'Sustentación del literal e):',
         'La comparacion se hizo contra una heuristica simple, como solicita el literal, y se reporto el error relativo posible. Como no hay optimos certificados, se aclara que el gap no es contra el optimo global sino contra la referencia experimental.')

doc.add_page_break()
# f
h = doc.add_heading('f) Analisis experimental de parametros', level=1)
doc.add_paragraph(
    'El analisis de parametros se realizo sobre el dataset con mas ciudades, cities_100_131415, porque representa el escenario mas exigente de los archivos disponibles. '
    'Se evaluaron variaciones en tamano de poblacion, tipo de crossover, probabilidad de crossover y probabilidad de mutacion. Para mantener un costo computacional razonable, el experimento uso 40 generaciones por configuracion y una mejora final 2-opt acotada.'
)
param_top = params[['pop_size', 'crossover', 'crossover_prob', 'mutation_prob', 'best_distance_ga_2opt', 'gap_vs_nn_pct', 'unique_pct_final', 'hamming_pct_final', 'valid_solution']].head(10).copy()
add_table_from_df(doc, param_top, ['pop_size', 'crossover', 'crossover_prob', 'mutation_prob', 'best_distance_ga_2opt', 'gap_vs_nn_pct', 'unique_pct_final', 'hamming_pct_final', 'valid_solution'],
                  ['Poblacion', 'Crossover', 'Pc', 'Pm', 'Distancia', 'Gap NN (%)', 'Unicos (%)', 'Hamming (%)', 'Valida'], font_size=7.0)
add_image(doc, FIG / 'parametros_top_configuraciones.png', 'Figura 6. Mejores configuraciones del analisis experimental de parametros.', width=6.6)

doc.add_paragraph(
    'El resultado no debe interpretarse como una regla universal, sino como evidencia experimental para estos datasets y este presupuesto de generaciones. En esta corrida, OX tuvo menor distancia promedio que PMX, la probabilidad de crossover 0.90 fue mejor que 0.80 y una mutacion de 0.20 mantuvo mejor equilibrio entre exploracion y explotacion que 0.05. '
    'El tamano de poblacion mostro un resultado interesante: una poblacion menor no fue necesariamente peor bajo un presupuesto corto de generaciones, porque poblaciones grandes pueden requerir mas generaciones para explotar adecuadamente la diversidad adicional.'
)
add_image(doc, FIG / 'parametros_poblacion_distancia.png', 'Figura 7. Efecto del tamano de poblacion en la distancia promedio.', width=5.8)
add_image(doc, FIG / 'parametros_mutacion_distancia.png', 'Figura 8. Efecto de la probabilidad de mutacion en la distancia promedio.', width=5.8)
add_image(doc, FIG / 'parametros_crossover_distancia.png', 'Figura 9. Efecto del tipo y probabilidad de crossover.', width=5.8)

param_summary = pd.DataFrame([
    ['Tamano de poblacion', '60, 100 y 140', 'A mayor poblacion hay mas diversidad potencial, pero tambien puede requerirse mas tiempo para convertirla en mejora.'],
    ['Tipo de crossover', 'OX y PMX', 'OX fue seleccionado como operador base porque preserva orden relativo y mantuvo buen rendimiento promedio.'],
    ['Probabilidad de crossover', '0.80 y 0.90', '0.90 favorecio mayor recombinacion de informacion entre padres y mejor resultado promedio en esta prueba.'],
    ['Probabilidad de mutacion', '0.05 y 0.20', '0.20 ayudo a sostener diversidad y evitar estancamiento temprano en el experimento.']
], columns=['Parametro', 'Valores evaluados', 'Interpretacion'])
add_table_from_df(doc, param_summary, ['Parametro', 'Valores evaluados', 'Interpretacion'], font_size=8.0)
add_note(doc, 'Sustentación del literal f):',
         'El analisis de parametros no solo reporta tablas; explica que la calidad del AG depende del equilibrio entre exploracion y explotacion. Por eso se compararon poblacion, crossover y mutacion, que son los factores pedidos en el literal.')

doc.add_page_break()
# Conclusions
h = doc.add_heading('Conclusiones tecnicas', level=1)
conclusions = [
    'La representacion por permutaciones es tecnicamente adecuada para el TSP porque convierte la factibilidad de la solucion en una propiedad directa del cromosoma: cada ciudad aparece una sola vez y no se requieren reparaciones posteriores para eliminar duplicados o completar faltantes.',
    'La matriz de distancias euclidianas permitio evaluar todas las rutas con el mismo criterio objetivo. Esto hizo posible comparar de manera consistente los datasets, la heuristica Nearest Neighbor, NN+2-opt y el AG elitista + 2-opt.',
    'El AG generacional elitista fue una eleccion justificada porque conserva las mejores rutas entre generaciones. En un problema combinatorio como el TSP, esta propiedad evita que una solucion de buena calidad desaparezca por azar durante el crossover o la mutacion.',
    'El elitismo no trabaja solo: se combino con seleccion por torneo para favorecer rutas prometedoras, OX para recombinar rutas validas y mutacion por inversion para introducir variacion sin generar ciudades repetidas ni faltantes.',
    'OX fue coherente con el problema porque es un operador de crossover para permutaciones. La decision metodologica no fue comparar OX contra el AG elitista, porque pertenecen a niveles distintos: el AG elitista es la arquitectura del algoritmo y OX es el operador interno que recombina rutas padres para producir hijos validos.',
    'La mutacion por inversion fue apropiada porque modifica el orden de visita de un segmento de la ruta sin alterar el conjunto de ciudades. Por eso mantiene la validez de la solucion y, al mismo tiempo, permite explorar nuevas configuraciones.',
    'El analisis del fitness cumplio el literal porque no solo observo el mejor individuo: tambien se graficaron fitness promedio y peor fitness. Esto permite analizar convergencia, separacion entre individuos y posibles senales de estancamiento.',
    'La diversidad poblacional se midio con porcentaje de individuos unicos y distancia Hamming promedio. Estas metricas permitieron distinguir una convergencia esperada de una posible convergencia prematura causada por perdida excesiva de variabilidad.',
    'La comparacion con Nearest Neighbor fue pertinente porque el literal solicitaba contrastar el AG con una heuristica simple. Adicionalmente, NN+2-opt ofrecio una referencia mas exigente para revisar si el AG era competitivo frente a una solucion localmente refinada.',
    'El uso de 2-opt como refinamiento final mejoro la ruta obtenida por el AG al corregir cruces innecesarios. Sin embargo, el nucleo de la propuesta sigue siendo evolutivo: 2-opt actua como mejora local posterior y no reemplaza al algoritmo genetico.',
    'El analisis experimental confirmo que el rendimiento depende del equilibrio entre tamano de poblacion, probabilidad de crossover y probabilidad de mutacion. Una poblacion mayor puede aportar diversidad, pero tambien puede requerir mas generaciones para aprovecharla.',
    'No se afirma optimalidad global porque los datasets no incluyen una ruta optima certificada. Por esta razon, el error relativo se reporto frente a referencias experimentales y no como brecha contra un optimo teorico desconocido.',
    'Como mejora futura, se recomienda ejecutar varias semillas por configuracion y reportar media, desviacion estandar y mejor valor por dataset. Esto fortaleceria la robustez estadistica del experimento.'
]
add_bullets(doc, conclusions)

doc.add_page_break()

# Participation
h = doc.add_heading('Resumen de participacion del Grupo 5', level=1)
participation_df = pd.DataFrame([
    ['Nancy Altamirano', 'Modelamiento del problema, datos y literal a', 'Revision de columnas city, x, y; verificacion de cantidad de ciudades; validacion de rutas sin ciudades repetidas/faltantes; apoyo en la explicacion de la representacion por permutaciones y en la preparacion de la tabla de validacion.'],
    ['Gustavo Berru', 'Implementacion del AG y operadores geneticos', 'Desarrollo de la poblacion inicial, seleccion por torneo, elitismo, funcion fitness, crossover OX y mutacion por inversion; verificacion de que los operadores generen descendientes validos; apoyo en los literales b y c.'],
    ['Raquel Pacheco', 'Evaluacion, comparacion y parametros', 'Ejecucion de Nearest Neighbor, NN+2-opt y AG+2-opt; calculo de gaps relativos; interpretacion de la calidad de solucion; diseno y analisis del experimento de parametros de poblacion, crossover y mutacion; apoyo en los literales e y f.'],
    ['Kevin Viteri', 'Diversidad, graficas y consolidacion tecnica', 'Medicion de diversidad con individuos unicos y distancia Hamming; analisis de convergencia y estancamiento; generacion e interpretacion de graficas de fitness, diversidad y parametros; consolidacion final del informe y apoyo en los literales c y d.']
], columns=['Integrante', 'Responsabilidad asignada', 'Actividades realizadas'])
add_table_from_df(doc, participation_df, ['Integrante', 'Responsabilidad asignada', 'Actividades realizadas'], font_size=7.7)

doc.add_paragraph(
    'La distribucion anterior mantiene una carga equilibrada: cada integrante participa en un componente tecnico verificable del proyecto y tambien puede sustentar una parte especifica durante la exposicion. La asignacion evita concentrar el trabajo en una sola persona y cubre datos, algoritmo, evaluacion, diversidad, parametros y documentacion final.'
)

# References
h = doc.add_heading('Referencias de apoyo', level=1)
refs = [
    'Goldberg, D. E. (1989). Genetic Algorithms in Search, Optimization, and Machine Learning. Addison-Wesley.',
    'Holland, J. H. (1975). Adaptation in Natural and Artificial Systems. University of Michigan Press.',
    'Mitchell, M. (1998). An Introduction to Genetic Algorithms. MIT Press.',
    'Oliver, I. M., Smith, D. J., & Holland, J. R. C. (1987). A study of permutation crossover operators on the traveling salesman problem. Proceedings of the Second International Conference on Genetic Algorithms.',
    'Croes, G. A. (1958). A method for solving traveling-salesman problems. Operations Research, 6(6), 791-812.'
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

# Appendix note
h = doc.add_heading('Archivos entregados en el paquete', level=1)
files_df = pd.DataFrame([
    ['code/tsp_ga_project.py', 'Codigo fuente reproducible del AG, heuristica NN, 2-opt, metricas y graficas.'],
    ['data/*.csv', 'Datasets originales usados en el proyecto.'],
    ['results/summary_all_datasets.csv', 'Resumen cuantitativo de calidad, diversidad, convergencia y validez.'],
    ['results/validacion_permutaciones.csv', 'Verificacion de que las rutas no tienen ciudades repetidas ni faltantes.'],
    ['results/parameter_experiment_largest_dataset.csv', 'Tabla completa del analisis experimental de parametros.'],
    ['figures/*.png', 'Graficas individuales y agregadas de fitness, diversidad, comparacion y parametros.']
], columns=['Archivo/carpeta', 'Contenido'])
add_table_from_df(doc, files_df, ['Archivo/carpeta', 'Contenido'], font_size=8.0)

out_docx = REPORT / 'informe_tsp_algoritmos_geneticos_grupo5_v4.docx'
doc.save(out_docx)
print(out_docx)
