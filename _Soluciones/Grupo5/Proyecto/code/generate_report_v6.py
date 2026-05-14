from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / 'results'
FIGURES = ROOT / 'figures'
REPORT = ROOT / 'report'
REPORT.mkdir(exist_ok=True)

summary = pd.read_csv(RESULTS / 'summary_all_datasets.csv')
validation = pd.read_csv(RESULTS / 'validacion_permutaciones.csv')
param = pd.read_csv(RESULTS / 'parameter_experiment_largest_dataset.csv')

# Valores globales para el informe
avg_nn = summary['nn_distance'].mean()
avg_ag2 = summary['ga_2opt_distance'].mean()
avg_gap = summary['gap_vs_nn_pct'].mean()
avg_gap_nn2 = summary['gap_vs_nn2_pct'].mean()
all_valid = bool(validation[['ruta_nn_valida','ruta_nn_2opt_valida','ruta_ag_valida','ruta_ag_2opt_valida','sin_ciudades_repetidas','sin_ciudades_faltantes']].all().all())
invalid_children_total = int(validation['hijos_invalidos_generados'].sum())
repaired_total = int(validation['hijos_reparados'].sum())
invalid_pop_total = int(validation['generaciones_con_poblacion_invalida'].sum())
stagnation_count = int(summary['stagnation'].str.contains('prematura', case=False, na=False).sum())
largest_dataset = summary.sort_values('n_ciudades', ascending=False).iloc[0]['dataset']

# Helpers de formato
BLUE = '1F4E79'
LIGHT_BLUE = 'D9EAF7'
LIGHT_GRAY = 'F2F2F2'
DARK_GRAY = '404040'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=8.5, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def fmt_num(v, decimals=2):
    if pd.isna(v):
        return '-'
    return f'{float(v):,.{decimals}f}'

def fmt_pct(v, decimals=2):
    if pd.isna(v):
        return '-'
    return f'{float(v):,.{decimals}f}%'

def bool_es(v):
    return 'Sí' if bool(v) else 'No'

# Crear documento
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.45)
sec.bottom_margin = Cm(1.45)
sec.left_margin = Cm(1.45)
sec.right_margin = Cm(1.45)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(10.5)
for h in ['Heading 1','Heading 2','Heading 3']:
    styles[h].font.name = 'Calibri'
    styles[h].font.color.rgb = RGBColor.from_string(BLUE)
styles['Heading 1'].font.size = Pt(17)
styles['Heading 2'].font.size = Pt(14)
styles['Heading 3'].font.size = Pt(12)

# Header/footer
section = doc.sections[0]
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.add_run('Grupo 5 - Resolución del TSP mediante Algoritmos Genéticos').font.size = Pt(8)

# Portada
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PROYECTO FINAL DE INTELIGENCIA ARTIFICIAL')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor.from_string(BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Resolución del TSP mediante Algoritmos Genéticos')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor.from_string(BLUE)

for text in ['Informe profesional mejorado y sustentado por literales', 'Grupo 5', 'Nancy Altamirano | Gustavo Berru | Raquel Pacheco | Kevin Viteri', 'Carrera: Ciencia de Datos', 'Mayo de 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    if text == 'Grupo 5':
        run.bold = True
        run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tema seleccionado: resolución del Problema del Viajante de Comercio (TSP) mediante un Algoritmo Genético Generacional Elitista con representación por permutaciones, crossover OX, mutación por inversión y refinamiento final 2-opt.')
run.italic = True
run.font.size = Pt(11)

doc.add_page_break()

# Índice manual
h = doc.add_heading('Índice del informe', level=1)
items = [
    'Resumen ejecutivo',
    'Objetivo, contexto y configuración del experimento',
    'a) Representación de individuos como permutaciones válidas',
    'b) Operadores genéticos seleccionados y justificación',
    'c) Análisis de la evolución del fitness',
    'd) Medición y análisis de diversidad poblacional',
    'e) Evaluación de la calidad de las soluciones',
    'f) Análisis experimental de parámetros',
    'Conclusiones técnicas',
    'Resumen de participación del Grupo 5',
    'Referencias de apoyo'
]
for item in items:
    doc.add_paragraph(item, style=None).style = doc.styles['Normal']

doc.add_page_break()

# Funciones para contenido
def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p

def add_bullets(items, level=0):
    for item in items:
        style = 'List Bullet' if level == 0 else 'List Bullet 2'
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, tuple):
            p.add_run(item[0]).bold = True
            p.add_run(item[1])
        else:
            p.add_run(str(item))

def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.add_run(str(item))

def add_table_from_df(df, columns, headers=None, decimals=None, font_size=8.0, title=None):
    if title:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
    headers = headers or columns
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for j, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], header, bold=True, size=font_size, color='000000')
        set_cell_shading(table.rows[0].cells[j], LIGHT_BLUE)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(columns):
            val = row[col]
            if isinstance(val, (float, np.floating)):
                dec = decimals.get(col, 2) if decimals else 2
                text = fmt_num(val, dec)
            elif isinstance(val, (bool, np.bool_)):
                text = bool_es(val)
            else:
                text = str(val)
            set_cell_text(cells[j], text, size=font_size)
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table

def add_figure(img_name, caption, width_cm=16.5):
    path = FIGURES / img_name
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    return cap

# Resumen ejecutivo
add_heading('Resumen ejecutivo', 1)
add_para('Este informe presenta la solución del Problema del Viajante de Comercio (TSP) mediante un Algoritmo Genético Generacional Elitista. El TSP requiere construir una ruta cerrada que visite cada ciudad exactamente una vez y retorne al punto inicial minimizando la distancia total. Debido a que el número de rutas posibles crece factorialmente con el número de ciudades, el uso de búsqueda exhaustiva resulta ineficiente para instancias medianas; por ello se aplica una técnica evolutiva capaz de explorar múltiples rutas candidatas y mejorar progresivamente la calidad de las soluciones.')
add_para(f'Se trabajaron los 10 datasets entregados, con tamaños entre {summary.n_ciudades.min()} y {summary.n_ciudades.max()} ciudades. No se agregaron ciudades artificiales, porque el objetivo fue respetar los datos reales suministrados. El algoritmo principal usó población de 100 individuos, 250 generaciones, selección por torneo, crossover OX, mutación por inversión, elitismo de dos individuos y refinamiento local final 2-opt.')
add_para(f'El resultado fue válido en todos los datasets: todas las rutas finales fueron permutaciones completas, sin ciudades repetidas ni faltantes. En comparación con la heurística Nearest Neighbor, la distancia promedio pasó de {fmt_num(avg_nn)} a {fmt_num(avg_ag2)}, lo que representa una mejora relativa promedio de {fmt_pct(avg_gap)} respecto a dicha heurística. Como el óptimo global certificado no fue proporcionado para los datasets, el error relativo se reporta contra la heurística de referencia, tal como permite el literal e).')
add_para('Idea clave para la presentación:', bold_start='Idea clave para la presentación:')
add_para('El AG elitista fue elegido como estructura principal porque conserva las mejores rutas encontradas y evita retrocesos por azar; OX no es un algoritmo alternativo al AG elitista, sino el operador de crossover usado dentro de ese algoritmo para preservar rutas válidas en forma de permutaciones.')

# Objetivo y configuración
add_heading('Objetivo, contexto y configuración del experimento', 1)
add_para('El objetivo del proyecto es resolver el TSP con un enfoque de Inteligencia Artificial basado en Algoritmos Genéticos y demostrar, de forma explícita, que se cumplen los literales a) hasta f) solicitados en el enunciado. Por ello, el informe está organizado en el mismo orden de los literales y cada sección incluye qué se hizo, por qué se hizo, cómo se verificó y qué evidencia se obtuvo.')
add_para('Configuración base utilizada en la ejecución principal:', bold_start='Configuración base utilizada en la ejecución principal:')
config_df = pd.DataFrame([
    ['Representación', 'Permutación de ciudades', 'Cada cromosoma contiene todos los índices de ciudad una sola vez.'],
    ['Población', '100 individuos', 'Permite diversidad inicial suficiente sin elevar demasiado el costo computacional.'],
    ['Generaciones', '250', 'Da tiempo para observar convergencia, estancamiento y evolución del fitness.'],
    ['Selección', 'Torneo k = 3', 'Mantiene presión selectiva moderada sin eliminar completamente la diversidad.'],
    ['Crossover', 'OX, probabilidad 0.90', 'Conserva orden relativo y evita rutas inválidas en problemas de permutación.'],
    ['Mutación', 'Inversión, probabilidad 0.15', 'Modifica segmentos de la ruta sin repetir ni eliminar ciudades.'],
    ['Elitismo', '2 mejores individuos', 'Protege las mejores rutas entre generaciones.'],
    ['Mejora local', '2-opt final', 'Refina la mejor ruta eliminando cruces o enlaces innecesarios.']
], columns=['Parámetro', 'Valor usado', 'Justificación'])
add_table_from_df(config_df, ['Parámetro','Valor usado','Justificación'], font_size=8.2, title='Tabla 1. Configuración base del algoritmo.')

# Datasets
add_heading('Datasets trabajados', 2)
datasets_df = summary[['dataset','n_ciudades']].copy()
add_table_from_df(datasets_df, ['dataset','n_ciudades'], headers=['Dataset','Cantidad de ciudades'], font_size=8.5, title='Tabla 2. Datasets procesados.')
add_para('Aunque los archivos tienen menos de 100 ciudades en algunos casos, se conservaron exactamente como fueron entregados. Esto evita alterar la dificultad real de cada instancia y permite que los resultados sean reproducibles con los CSV originales.')

# Literal a
add_heading('a) Representación de cada individuo como una permutación válida de ciudades', 1)
add_para('Qué pide el literal:', bold_start='Qué pide el literal:')
add_bullets([
    'Representar cada individuo como una permutación válida de ciudades.',
    'Garantizar que las soluciones generadas en la inicialización, el crossover y la mutación sean válidas.',
    'Demostrar explícitamente que las rutas no tienen ciudades repetidas.',
    'Demostrar explícitamente que las rutas no tienen ciudades faltantes.'
])
add_para('Respuesta implementada:', bold_start='Respuesta implementada:')
add_para('Cada individuo se representó como un vector ordenado de índices de ciudades. Si un dataset tiene n ciudades, un cromosoma válido contiene exactamente los valores 0, 1, 2, ..., n-1. La posición del índice dentro del vector define el orden de visita; por tanto, el cromosoma no almacena coordenadas, sino la secuencia de ciudades que construye la ruta.')
add_para('Ejemplo conceptual: [4, 1, 7, 0, 3, 2, 6, 5] significa que la ruta inicia en la ciudad 4, luego visita la ciudad 1, después la ciudad 7, y así sucesivamente hasta volver a la ciudad inicial. Esta representación es correcta para el TSP porque el problema no consiste en decidir si una ciudad se visita o no, sino en decidir en qué orden se visitan todas las ciudades.')
add_para('Garantía de validez por etapa:', bold_start='Garantía de validez por etapa:')
valid_stage = pd.DataFrame([
    ['Inicialización', 'Se genera cada cromosoma con una permutación aleatoria de 0 a n-1. También se incluye una ruta semilla de Nearest Neighbor para iniciar con una solución razonable.', 'No puede haber repetidas ni faltantes porque la función de permutación genera cada índice una sola vez.'],
    ['Crossover', 'Se aplica Order Crossover (OX). Copia un segmento de un padre y completa los espacios con ciudades del segundo padre en el orden en que aparecen.', 'OX evita duplicados porque no inserta genes ya usados y completa con los faltantes.'],
    ['Mutación', 'Se aplica mutación por inversión. Selecciona dos posiciones e invierte el segmento entre ellas.', 'La inversión solo cambia el orden; no crea ni elimina ciudades.'],
    ['Validación defensiva', 'El código verifica cada ruta con is_valid_permutation y tiene una reparación defensiva si apareciera una ruta inválida.', 'En la ejecución real no se activó reparación porque OX + inversión conservaron la validez.']
], columns=['Etapa', 'Mecanismo aplicado', 'Cómo garantiza validez'])
add_table_from_df(valid_stage, ['Etapa','Mecanismo aplicado','Cómo garantiza validez'], font_size=7.8, title='Tabla 3. Control de validez en inicialización, crossover y mutación.')
val_df = validation[['dataset','n_ciudades','sin_ciudades_repetidas','sin_ciudades_faltantes','hijos_invalidos_generados','hijos_reparados','generaciones_con_poblacion_invalida']].copy()
add_table_from_df(val_df, ['dataset','n_ciudades','sin_ciudades_repetidas','sin_ciudades_faltantes','hijos_invalidos_generados','hijos_reparados','generaciones_con_poblacion_invalida'], headers=['Dataset','n','Sin repetidas','Sin faltantes','Hijos inválidos','Reparados','Poblaciones inválidas'], font_size=7.4, title='Tabla 4. Evidencia de rutas válidas por dataset.')
add_para(f'Sustentación del literal a): la validez no se dejó al azar. La estructura de permutación, el crossover OX y la mutación por inversión fueron seleccionados porque respetan la naturaleza combinatoria del TSP. La tabla anterior confirma que todos los datasets terminaron con rutas válidas; además, se registraron {invalid_children_total} hijos inválidos, {repaired_total} reparaciones y {invalid_pop_total} generaciones con población inválida. Por tanto, se cumple explícitamente la condición de no repetir ciudades y no omitir ciudades.')

# Literal b
add_heading('b) Definición de operadores genéticos y justificación del uso', 1)
add_para('Qué pide el literal:', bold_start='Qué pide el literal:')
add_bullets([
    'Definir un operador de crossover válido para permutaciones, como PMX, OX o CX.',
    'Definir un operador de mutación válido, como swap, inversión o scramble.',
    'Justificar por qué esos operadores son adecuados para el TSP.'
])
add_para('Respuesta implementada:', bold_start='Respuesta implementada:')
add_para('El algoritmo elegido fue un Algoritmo Genético Generacional Elitista. Dentro de ese algoritmo se utilizaron selección por torneo, crossover OX, mutación por inversión y elitismo. Es importante diferenciar la variante del algoritmo de sus operadores: el elitismo define cómo se conserva la mejor solución entre generaciones; OX define cómo se cruzan dos rutas padres para formar un hijo válido.')
add_para('Análisis conceptual:', bold_start='Análisis conceptual:')
add_para('La elección no fue entre OX y el AG elitista, porque pertenecen a niveles diferentes de la solución. El AG generacional elitista es el enfoque principal de búsqueda; OX es el operador de crossover usado dentro de ese AG para recombinar rutas sin romper la estructura de permutación. Por tanto, la formulación correcta es: “Algoritmo Genético Generacional Elitista con crossover OX y mutación por inversión”.')
operators = pd.DataFrame([
    ['Selección por torneo', 'Se eligen k individuos al azar y se reproduce el de menor distancia.', 'Controla la presión selectiva y permite que las mejores rutas tengan más probabilidad de reproducirse sin eliminar toda la diversidad.'],
    ['Crossover OX', 'Copia un segmento de un padre y completa con el orden relativo del otro padre.', 'Es válido para permutaciones: evita ciudades repetidas y faltantes, y conserva información de orden, que es crítica en el TSP.'],
    ['Mutación por inversión', 'Invierte un segmento de la ruta.', 'Mantiene todos los genes y solo cambia el orden de visita; además, es coherente con mejoras típicas del TSP.'],
    ['Elitismo', 'Copia los mejores individuos directamente a la siguiente generación.', 'Evita que una buena ruta se pierda por azar durante crossover o mutación.'],
    ['2-opt final', 'Invierte segmentos si reduce la distancia total.', 'Refina la mejor ruta hallada por el AG; por eso el enfoque puede considerarse híbrido o memético en la etapa final.']
], columns=['Operador o componente', 'Funcionamiento', 'Justificación'])
add_table_from_df(operators, ['Operador o componente','Funcionamiento','Justificación'], font_size=7.8, title='Tabla 5. Operadores genéticos y justificación técnica.')
add_para('Sustentación del literal b): OX fue seleccionado porque el TSP es sensible al orden de visita. Un cruce tradicional podría producir cromosomas inválidos, por ejemplo con ciudades repetidas y otras ausentes. OX reduce ese riesgo porque preserva un segmento de una ruta padre y completa la solución con las ciudades que faltan siguiendo el orden del otro padre. La mutación por inversión se seleccionó porque explora cambios locales sin destruir la estructura de permutación. El elitismo se incorporó para proteger la mejor solución alcanzada y estabilizar la convergencia.')

# Literal c
add_heading('c) Análisis de la evolución del fitness', 1)
add_para('Qué pide el literal:', bold_start='Qué pide el literal:')
add_bullets([
    'Graficar por generación el mejor fitness, el fitness promedio y el peor fitness.',
    'Analizar la convergencia del algoritmo.',
    'Analizar la presencia de estancamiento o convergencia prematura.'
])
add_para('Definición de fitness:', bold_start='Definición de fitness:')
add_para('El objetivo del TSP es minimizar la distancia total. Para adaptar el problema al lenguaje de fitness, se definió fitness = 1 / distancia_total. Con esta definición, una ruta más corta tiene mayor fitness. En cada generación se registraron tres métricas: mejor fitness, fitness promedio y peor fitness. Estas tres curvas permiten observar si la población mejora, si se concentra alrededor de buenas soluciones o si todavía mantiene individuos de baja calidad.')
add_figure('panel_fitness_all_datasets_v5.png', 'Figura 1. Evolución por generación del mejor fitness, fitness promedio y peor fitness para los 10 datasets.', width_cm=16.7)
add_para('Análisis de la gráfica de fitness:', bold_start='Análisis de la gráfica de fitness:')
add_para('En los 10 datasets se observa el comportamiento esperado de un algoritmo genético: el mejor fitness crece de manera rápida al inicio, porque la población encuentra rutas claramente mejores que las aleatorias; después, las mejoras se vuelven más pequeñas y la curva se estabiliza. El fitness promedio tiende a acercarse al mejor fitness, lo que evidencia convergencia de la población hacia rutas de mayor calidad. El peor fitness también mejora o se estabiliza porque la selección y el reemplazo generacional reducen la presencia de rutas de mala calidad.')
add_figure('panel_convergencia_distancia_all_datasets_v5.png', 'Figura 2. Convergencia por distancia: mejor distancia acumulada, distancia promedio y peor distancia por generación.', width_cm=16.7)
conv_df = summary[['dataset','best_generation','final_no_improve','unique_pct_final','hamming_pct_final','stagnation']].copy()
add_table_from_df(conv_df, ['dataset','best_generation','final_no_improve','unique_pct_final','hamming_pct_final','stagnation'], headers=['Dataset','Mejor generación','Gen. finales sin mejora','Únicos finales (%)','Hamming final (%)','Diagnóstico'], decimals={'unique_pct_final':1,'hamming_pct_final':1}, font_size=6.8, title='Tabla 6. Evidencia de convergencia y estancamiento por dataset.')
add_para(f'Análisis de convergencia y estancamiento:', bold_start='Análisis de convergencia y estancamiento:')
add_para(f'El criterio de convergencia no se basó únicamente en observar las curvas. También se registró la generación en la que apareció la mejor solución, la cantidad de generaciones finales sin mejora y la diversidad final. En {10 - stagnation_count} de los 10 datasets se clasificó la convergencia como activa o con estancamiento parcial controlado, porque todavía existieron mejoras o diversidad residual. En {stagnation_count} datasets se identificó posible convergencia prematura, debido a un tramo final largo sin mejora y diversidad final baja. Esto no invalida el algoritmo; más bien muestra que, en esas instancias, sería conveniente aumentar mutación, población o generaciones si se quisiera seguir explorando.')
add_para('Sustentación del literal c): se graficaron las tres curvas de fitness solicitadas para cada dataset y se complementaron con curvas de distancia para interpretar la convergencia. Además, el estancamiento se analizó con indicadores cuantitativos y no solo con apreciación visual.')

# Literal d
add_heading('d) Medición y análisis de diversidad poblacional', 1)
add_para('Qué pide el literal:', bold_start='Qué pide el literal:')
add_bullets([
    'Definir una métrica de diversidad, como distancia promedio entre individuos y/o porcentaje de individuos únicos.',
    'Graficar la diversidad a lo largo de las generaciones.',
    'Analizar la relación entre diversidad y calidad de solución.',
    'Analizar la relación entre diversidad y convergencia.'
])
add_para('Métricas de diversidad usadas:', bold_start='Métricas de diversidad usadas:')
add_bullets([
    ('Porcentaje de individuos únicos: ', 'mide qué proporción de la población contiene rutas distintas. Si baja demasiado, la población se vuelve homogénea.'),
    ('Distancia Hamming promedio normalizada: ', 'mide el porcentaje promedio de posiciones diferentes entre pares de rutas. Sirve como distancia promedio entre individuos en una representación por permutaciones.')
])
add_figure('panel_diversidad_all_datasets_v5.png', 'Figura 3. Evolución de diversidad poblacional por generación: individuos únicos y distancia Hamming promedio.', width_cm=16.7)
add_para('Análisis de la diversidad a lo largo de las generaciones:', bold_start='Análisis de la diversidad a lo largo de las generaciones:')
add_para('La diversidad disminuye conforme avanza el algoritmo porque la selección y el elitismo favorecen la reproducción de rutas de mejor calidad. Esta reducción es normal: el algoritmo empieza explorando muchas rutas distintas y luego explota las regiones más prometedoras. Sin embargo, si la diversidad cae demasiado pronto, puede aparecer convergencia prematura, porque la población deja de explorar rutas alternativas.')
add_figure('diversidad_final_por_dataset_v5.png', 'Figura 4. Porcentaje final de individuos únicos por dataset.', width_cm=16.2)
add_figure('diversidad_vs_calidad_solucion.png', 'Figura 5. Relación entre diversidad final y calidad de solución medida como gap relativo frente a Nearest Neighbor.', width_cm=15.0)
add_figure('diversidad_vs_convergencia.png', 'Figura 6. Relación entre diversidad final y convergencia/estancamiento medida por generaciones finales sin mejora.', width_cm=15.0)
rel_df = summary[['dataset','corr_unique_vs_best_distance','corr_hamming_vs_best_distance','corr_unique_vs_no_improve']].copy()
add_table_from_df(rel_df, ['dataset','corr_unique_vs_best_distance','corr_hamming_vs_best_distance','corr_unique_vs_no_improve'], headers=['Dataset','Corr. únicos vs distancia','Corr. Hamming vs distancia','Corr. únicos vs no mejora'], decimals={'corr_unique_vs_best_distance':3,'corr_hamming_vs_best_distance':3,'corr_unique_vs_no_improve':3}, font_size=7.5, title='Tabla 7. Relación cuantitativa entre diversidad, calidad y convergencia.')
add_para('Análisis diversidad-calidad:', bold_start='Análisis diversidad-calidad:')
add_para('La relación no debe interpretarse como “más diversidad siempre es mejor”. Al inicio, una diversidad alta es necesaria para explorar rutas diferentes. En etapas finales, una diversidad moderadamente menor puede ser positiva porque indica que la población se concentra en buenas soluciones. El riesgo aparece cuando la diversidad cae demasiado temprano y el algoritmo ya no mejora. En los resultados, los gaps negativos frente a Nearest Neighbor muestran que el algoritmo obtuvo soluciones de buena calidad aun cuando la diversidad final bajó en varios datasets.')
add_para('Análisis diversidad-convergencia:', bold_start='Análisis diversidad-convergencia:')
add_para('Cuando la diversidad final es baja y el número de generaciones sin mejora es alto, existe evidencia de estancamiento. Esto ocurrió principalmente en los datasets señalados como posible convergencia prematura. En términos prácticos, esta información permite defender decisiones de ajuste: aumentar la probabilidad de mutación, ampliar la población o prolongar generaciones puede reintroducir exploración si la diversidad cae demasiado pronto.')
add_para('Sustentación del literal d): se definieron dos métricas de diversidad, se graficaron a lo largo de las generaciones y se analizó su relación con calidad y convergencia mediante gráficas y correlaciones. Por tanto, el literal no queda solo en una definición, sino que se evalúa con evidencia experimental.')

# Literal e
add_heading('e) Evaluación de la calidad de las soluciones', 1)
add_para('Qué pide el literal:', bold_start='Qué pide el literal:')
add_bullets([
    'Comparar los resultados del AG con al menos una heurística simple, por ejemplo Nearest Neighbor.',
    'Reportar fitness final y error relativo cuando sea posible.'
])
add_para('Heurística de comparación:', bold_start='Heurística de comparación:')
add_para('Se utilizó Nearest Neighbor como heurística simple. Esta heurística inicia en una ciudad y, en cada paso, elige la ciudad no visitada más cercana. Para hacer una comparación más justa, se probó cada ciudad como posible inicio y se conservó la mejor ruta obtenida. Además, se calculó una versión Nearest Neighbor + 2-opt para comparar el AG contra una heurística refinada.')
add_figure('comparacion_nn_vs_ag2opt_v5.png', 'Figura 7. Comparación de distancia final: Nearest Neighbor vs AG elitista + 2-opt.', width_cm=16.7)
add_figure('gap_relativo_vs_nn_v5.png', 'Figura 8. Error relativo frente a Nearest Neighbor. Valores negativos representan mejora del AG.', width_cm=16.2)
quality_df = summary[['dataset','n_ciudades','nn_distance','nn_2opt_distance','ga_2opt_distance','gap_vs_nn_pct','gap_vs_nn2_pct','valid_solution']].copy()
quality_df['fitness_final'] = 1 / quality_df['ga_2opt_distance']
quality_df = quality_df[['dataset','n_ciudades','nn_distance','nn_2opt_distance','ga_2opt_distance','fitness_final','gap_vs_nn_pct','gap_vs_nn2_pct','valid_solution']]
add_table_from_df(quality_df, ['dataset','n_ciudades','nn_distance','nn_2opt_distance','ga_2opt_distance','fitness_final','gap_vs_nn_pct','gap_vs_nn2_pct','valid_solution'], headers=['Dataset','n','NN dist.','NN+2opt dist.','AG+2opt dist.','Fitness final','Error vs NN','Error vs NN+2opt','Ruta válida'], decimals={'nn_distance':2,'nn_2opt_distance':2,'ga_2opt_distance':2,'fitness_final':6,'gap_vs_nn_pct':2,'gap_vs_nn2_pct':2}, font_size=6.7, title='Tabla 8. Calidad de solución, fitness final y error relativo.')
add_para(f'Análisis de calidad:', bold_start='Análisis de calidad:')
add_para(f'El AG elitista + 2-opt superó a Nearest Neighbor en todos los datasets, porque los errores relativos frente a NN fueron negativos. En promedio, la distancia se redujo en {fmt_pct(avg_gap)} respecto a NN. Frente a NN+2-opt, el promedio fue {fmt_pct(avg_gap_nn2)}, lo que indica que el AG también fue competitivo frente a una heurística simple refinada, aunque en algunos datasets NN+2-opt quedó ligeramente por debajo. Esto es esperable porque 2-opt es una mejora local fuerte y el AG no garantiza óptimo global, sino buenas soluciones aproximadas.')
add_para('Nota sobre el error relativo:', bold_start='Nota sobre el error relativo:')
add_para('Como los datasets no incluyen una solución óptima certificada, no sería correcto reportar error relativo contra el óptimo global. Por transparencia, el error relativo se reportó respecto a la heurística Nearest Neighbor y adicionalmente respecto a Nearest Neighbor + 2-opt. Esta decisión cumple el literal e) “cuando sea posible” y evita afirmar una precisión que no fue verificada.')
add_para('Sustentación del literal e): se comparó el AG contra una heurística simple, se reportó el fitness final y se calculó error relativo contra la referencia disponible. Las rutas finales fueron válidas en todos los casos, por lo que la comparación no mezcla soluciones inválidas ni incompletas.')

# Literal f
add_heading('f) Análisis experimental de parámetros', 1)
add_para('Qué pide el literal:', bold_start='Qué pide el literal:')
add_bullets([
    'Reportar resultados en tablas.',
    'Analizar el tamaño de población.',
    'Analizar la parametrización del crossover.',
    'Analizar la probabilidad de mutación.'
])
add_para('Diseño del experimento:', bold_start='Diseño del experimento:')
add_para(f'El análisis de parámetros se realizó sobre el dataset con mayor número de ciudades: {largest_dataset}. Se eligió ese archivo porque representa el escenario más exigente entre los datos disponibles. Para mantener un costo computacional razonable, cada configuración se evaluó con 40 generaciones y luego se aplicó 2-opt final. Las variables analizadas fueron tamaño de población, tipo de crossover, probabilidad de crossover y probabilidad de mutación.')
param_design = pd.DataFrame([
    ['Tamaño de población', '60, 100, 140', 'Permite observar si más individuos mejoran exploración y calidad.'],
    ['Tipo de crossover', 'OX y PMX', 'Ambos son válidos para permutaciones; se compara su desempeño experimental.'],
    ['Probabilidad de crossover', '0.80 y 0.90', 'Mide cuánto influye recombinar padres frente a copiar un padre.'],
    ['Probabilidad de mutación', '0.05 y 0.20', 'Evalúa el balance entre exploración adicional y perturbación excesiva.'],
    ['Generaciones por configuración', '40', 'Presupuesto uniforme para comparar parámetros bajo la misma restricción.']
], columns=['Parámetro analizado','Valores probados','Razón de análisis'])
add_table_from_df(param_design, ['Parámetro analizado','Valores probados','Razón de análisis'], font_size=8.0, title='Tabla 9. Diseño del análisis experimental de parámetros.')
top_param = param.head(10).copy()
add_table_from_df(top_param, ['pop_size','crossover','crossover_prob','mutation_prob','best_distance_ga_2opt','gap_vs_nn_pct','unique_pct_final','hamming_pct_final'], headers=['Población','Crossover','Pc','Pm','Dist. AG+2opt','Error vs NN','Únicos finales (%)','Hamming final (%)'], decimals={'crossover_prob':2,'mutation_prob':2,'best_distance_ga_2opt':2,'gap_vs_nn_pct':2,'unique_pct_final':1,'hamming_pct_final':1}, font_size=7.4, title='Tabla 10. Diez mejores configuraciones del experimento de parámetros.')
add_figure('parametros_top_configuraciones.png', 'Figura 9. Mejores configuraciones del análisis experimental de parámetros.', width_cm=16.7)
add_figure('parametros_poblacion_distancia.png', 'Figura 10. Efecto del tamaño de población sobre la distancia promedio.', width_cm=14.5)
add_figure('parametros_mutacion_distancia.png', 'Figura 11. Efecto de la probabilidad de mutación sobre la distancia promedio.', width_cm=14.5)
add_figure('parametros_crossover_distancia.png', 'Figura 12. Efecto del tipo y probabilidad de crossover sobre la distancia promedio.', width_cm=14.5)
# Tablas agregadas por parámetro
pop_summary = param.groupby('pop_size', as_index=False).agg(distancia_promedio=('best_distance_ga_2opt','mean'), diversidad_promedio=('unique_pct_final','mean'), hamming_promedio=('hamming_pct_final','mean'))
mut_summary = param.groupby('mutation_prob', as_index=False).agg(distancia_promedio=('best_distance_ga_2opt','mean'), diversidad_promedio=('unique_pct_final','mean'), hamming_promedio=('hamming_pct_final','mean'))
cx_summary = param.groupby(['crossover','crossover_prob'], as_index=False).agg(distancia_promedio=('best_distance_ga_2opt','mean'), diversidad_promedio=('unique_pct_final','mean'))
add_table_from_df(pop_summary, ['pop_size','distancia_promedio','diversidad_promedio','hamming_promedio'], headers=['Población','Distancia promedio','Únicos prom. (%)','Hamming prom. (%)'], decimals={'distancia_promedio':2,'diversidad_promedio':1,'hamming_promedio':1}, font_size=8.0, title='Tabla 11. Resumen por tamaño de población.')
add_table_from_df(mut_summary, ['mutation_prob','distancia_promedio','diversidad_promedio','hamming_promedio'], headers=['Prob. mutación','Distancia promedio','Únicos prom. (%)','Hamming prom. (%)'], decimals={'mutation_prob':2,'distancia_promedio':2,'diversidad_promedio':1,'hamming_promedio':1}, font_size=8.0, title='Tabla 12. Resumen por probabilidad de mutación.')
add_table_from_df(cx_summary, ['crossover','crossover_prob','distancia_promedio','diversidad_promedio'], headers=['Crossover','Pc','Distancia promedio','Únicos prom. (%)'], decimals={'crossover_prob':2,'distancia_promedio':2,'diversidad_promedio':1}, font_size=8.0, title='Tabla 13. Resumen por parametrización del crossover.')
add_para('Análisis del tamaño de población:', bold_start='Análisis del tamaño de población:')
add_para('Una población mayor tiende a incrementar la exploración porque mantiene más rutas candidatas por generación. Sin embargo, también aumenta el costo computacional. En el experimento, el tamaño poblacional no se interpreta de forma aislada, sino junto con mutación y crossover. Esto es importante para la defensa: una población de 140 puede explorar más, pero si la mutación es baja o el crossover no introduce suficiente variación, la mejora no será necesariamente proporcional.')
add_para('Análisis del crossover:', bold_start='Análisis del crossover:')
add_para('OX y PMX son operadores válidos para permutaciones. OX fue elegido en la ejecución principal porque conserva el orden relativo de visita, que es una característica central en el TSP. El experimento incluyó PMX para demostrar que existen alternativas válidas y para comparar su efecto bajo el mismo presupuesto de generaciones. La probabilidad de crossover alta favorece la recombinación de rutas, pero si se usa sin mutación suficiente puede acelerar la homogeneización de la población.')
add_para('Análisis de la mutación:', bold_start='Análisis de la mutación:')
add_para('La mutación por inversión introduce exploración sin romper la validez de la ruta. Una probabilidad demasiado baja puede causar convergencia prematura porque la población deja de explorar. Una probabilidad demasiado alta puede dificultar la explotación de buenas rutas. En este experimento, 0.20 mostró configuraciones competitivas porque permitió recuperar diversidad en un problema donde el orden de visita es muy sensible.')
add_para('Sustentación del literal f): el análisis no se limitó a mencionar parámetros; se probaron configuraciones, se reportaron tablas y se graficó el efecto de población, crossover y mutación. Esto permite defender por qué la configuración base es razonable y cómo podrían ajustarse los parámetros si se detecta estancamiento.')

# Conclusiones
add_heading('Conclusiones técnicas', 1)
conclusions = [
    'La representación por permutación es la más adecuada para el TSP porque cada ruta debe contener todas las ciudades exactamente una vez. Esta representación evita modelar el problema como una cadena binaria, que sería menos natural para rutas ordenadas.',
    'La validez de las rutas se garantizó en todas las etapas: inicialización mediante permutaciones, crossover OX mediante preservación y completado sin duplicados, y mutación por inversión mediante reordenamiento sin pérdida de genes.',
    'El AG elitista fue una elección pertinente porque protege las mejores soluciones encontradas. Esto es especialmente útil en el TSP, donde una buena ruta puede perderse por operaciones aleatorias si no existe elitismo.',
    'OX no compite con el AG elitista; OX es el operador de cruce dentro del AG. La estructura principal es el AG generacional elitista, mientras que OX permite recombinar rutas sin violar la restricción de permutación.',
    'El fitness definido como inverso de la distancia permitió transformar un problema de minimización en una medición de maximización, lo que facilita la interpretación evolutiva del algoritmo.',
    'Las gráficas de fitness y distancia evidencian mejora inicial rápida seguida de estabilización, comportamiento típico de algoritmos genéticos en problemas combinatorios.',
    'La diversidad poblacional disminuyó conforme la población convergió hacia mejores rutas. Esta reducción es normal, pero debe vigilarse porque puede producir convergencia prematura en algunos datasets.',
    'La comparación contra Nearest Neighbor demuestra que el AG elitista + 2-opt obtuvo mejores distancias en todos los datasets respecto a la heurística simple.',
    'El uso de 2-opt como refinamiento final fortaleció la calidad de las rutas al corregir segmentos locales ineficientes; por ello, el enfoque puede describirse como AG elitista con mejora local final.',
    'El análisis de parámetros mostró que la calidad final depende del equilibrio entre exploración y explotación. Población, crossover y mutación deben seleccionarse conjuntamente, no de forma aislada.',
    'El proyecto cumple los literales a) a f) porque incluye representación válida, operadores justificados, gráficas de fitness, métricas de diversidad, comparación con heurística y experimentación de parámetros con tablas y gráficas.',
    'Una posible mejora futura sería ejecutar múltiples semillas por configuración para estimar variabilidad estadística, así como comparar con otras heurísticas como inserción más cercana, simulated annealing o colonia de hormigas.'
]
add_numbered(conclusions)

# Participación
add_heading('Resumen de participación del Grupo 5', 1)
participation = pd.DataFrame([
    ['Nancy Altamirano', 'Preparación de datos, representación y validación', 'Organizó los CSV, verificó columnas city, x, y, controló cantidad de ciudades por dataset, apoyó el literal a), validó rutas sin ciudades repetidas ni faltantes y revisó la consistencia de las tablas de validación.'],
    ['Gustavo Berru', 'Diseño e implementación del algoritmo genético', 'Implementó la estructura del AG generacional elitista, selección por torneo, elitismo, crossover OX, mutación por inversión, validaciones de permutación, registro de histórico por generación y exportación de rutas finales.'],
    ['Raquel Pacheco', 'Evaluación de calidad y análisis experimental', 'Implementó y revisó la comparación contra Nearest Neighbor, cálculo de fitness final, error relativo, análisis del dataset más grande, experimentos de población, crossover y mutación, y consolidación de resultados cuantitativos.'],
    ['Kevin Viteri', 'Análisis gráfico, diversidad y documentación final', 'Generó las gráficas de fitness, convergencia, diversidad, comparación y parámetros; analizó estancamiento y convergencia prematura; integró la documentación profesional y revisó ortografía, estructura y sustentación para exposición.']
], columns=['Integrante','Responsabilidad asignada','Actividades realizadas'])
add_table_from_df(participation, ['Integrante','Responsabilidad asignada','Actividades realizadas'], font_size=7.8, title='Tabla 14. Participación equilibrada de los integrantes del Grupo 5.')
add_para('La distribución se planteó de manera balanceada para que cada integrante participe en una parte técnica del proyecto: datos y representación, algoritmo, evaluación experimental, y análisis/documentación. Esto permite defender el trabajo como un desarrollo grupal integral y no como una división basada en tareas menores.')

# Referencias
add_heading('Referencias de apoyo', 1)
refs = [
    'Croes, G. A. (1958). A method for solving traveling-salesman problems. Operations Research, 6(6), 791-812.',
    'Goldberg, D. E. (1989). Genetic algorithms in search, optimization, and machine learning. Addison-Wesley.',
    'Holland, J. H. (1975). Adaptation in natural and artificial systems. University of Michigan Press.',
    'Reinelt, G. (1994). The traveling salesman: Computational solutions for TSP applications. Springer.',
    'Whitley, D. (1994). A genetic algorithm tutorial. Statistics and Computing, 4, 65-85.'
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)

# Guardar
out = REPORT / 'informe_tsp_algoritmos_geneticos_grupo5_v6.docx'
doc.save(out)
print(out)
